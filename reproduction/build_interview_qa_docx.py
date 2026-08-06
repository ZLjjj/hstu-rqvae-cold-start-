from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCES = (
    ROOT / "RQVAE_HSTU_INTERVIEW_100_QA.md",
    ROOT / "RQVAE_HSTU_INTERVIEW_FULL_TOKEN_50_QA.md",
)
OUTPUT = ROOT / "RQ-VAE_HSTU_冷启动推荐项目_150个面试问题与参考答案.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 100, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
ZH_FONT = "Hiragino Sans GB"
EN_FONT = "Arial"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=EN_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), ZH_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_inline_runs(paragraph, text, base_size=11, base_color=None):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size - 0.5, color=DARK_BLUE, name="Menlo")
            run.font.highlight_color = None
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color)


def parse_source():
    sections = []
    for source in SOURCES:
        current_section = None
        current_question = None
        for line in source.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if stripped.startswith("## "):
                current_section = {"title": stripped[3:], "questions": []}
                sections.append(current_section)
                current_question = None
            elif stripped.startswith("### "):
                if current_section is None:
                    continue
                current_question = {"question": stripped[4:], "answer": []}
                current_section["questions"].append(current_question)
            elif stripped.startswith("参考答案：") and current_question is not None:
                current_question["answer"].append(stripped[len("参考答案："):])
            elif stripped and current_question is not None:
                current_question["answer"].append(stripped)
    return sections


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(32, 38, 45)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = EN_FONT
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = EN_FONT
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = EN_FONT
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = DARK_BLUE
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("RQ-VAE + HSTU 冷启动推荐｜面试准备手册")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(p)
    run = p.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("推荐算法项目面试手册")
    set_run_font(run, size=12, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RQ-VAE + HSTU 冷启动推荐")
    set_run_font(run, size=28, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("150 个面试问题与参考答案")
    set_run_font(run, size=17, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("覆盖：数据协议｜RQ-VAE｜HSTU｜Route A/B｜Full-Token｜指标｜工程复现｜实验批判")
    set_run_font(run, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 MovieLens-1M 严格物品冷启动复现实验")
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    doc.add_page_break()


def add_front_matter(doc, sections):
    p = doc.add_paragraph("使用说明", style="Heading 1")
    p.paragraph_format.page_break_before = False
    body = doc.add_paragraph()
    add_inline_runs(body, "建议采用三轮复习法：第一轮掌握项目叙事和关键数字；第二轮理解公式、代码路径与实验协议；第三轮针对缺陷和改进进行压力追问。回答时先给结论，再解释机制，最后补充数据或局限。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    add_inline_runs(p, "优先题目：1、5、18、49、63、68、77、78、91、101、102、121、128、131、135、141、146、149、150。面试中不要回避验证/测试重复、冷物品负采样和 Full-Token 冷指标下降，主动说明证据与结论边界。", base_color=DARK_BLUE)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    pPr.append(shd)

    doc.add_paragraph("关键实验数字", style="Heading 2")
    rows = [
        ("方法", "全体 AUC", "HR@100", "cold AUC", "cold HR@100"),
        ("HSTU Baseline", "0.8240", "0.5831", "0.0882", "0.0000"),
        ("Dense MiniLM", "0.8472", "0.4613", "0.4757", "0.0202"),
        ("Route A", "0.8575", "0.4505", "0.5920", "0.1028"),
        ("旧 Route B", "0.7485", "0.2636", "0.2289", "0.0018"),
        ("Full-Token", "0.7885", "0.3788", "0.0777", "0.0018"),
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(
                run,
                size=9.5,
                color=WHITE if i == 0 else None,
                bold=(
                    i == 0
                    or (i == 3 and j in (0, 3, 4))
                    or (i == 5 and j in (0, 2))
                ),
            )
            if i == 0:
                set_cell_shading(cell, "2E74B5")
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2500, 1600, 1700, 1750, 1810])

    doc.add_paragraph("模块索引", style="Heading 2")
    table = doc.add_table(rows=len(sections) + 1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0]
    for j, value in enumerate(("题号", "模块")):
        header.cells[j].text = ""
        p = header.cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=10, color=WHITE, bold=True)
        set_cell_shading(header.cells[j], "2E74B5")
    for i, section in enumerate(sections, start=1):
        match = re.search(r"（(\d+)—(\d+)）", section["title"])
        range_text = f"{match.group(1)}—{match.group(2)}" if match else ""
        title = re.sub(r"（\d+—\d+）", "", section["title"])
        for j, value in enumerate((range_text, title)):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=10)
            if i % 2 == 0:
                set_cell_shading(cell, "F4F6F9")
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1700, 7660])


def add_qa_sections(doc, sections):
    for section_idx, section in enumerate(sections):
        title = doc.add_paragraph(section["title"], style="Heading 1")
        title.paragraph_format.page_break_before = (
            section_idx > 0 and section_idx != len(sections) - 1
        )
        for item in section["questions"]:
            q = doc.add_paragraph(item["question"], style="Heading 2")
            q.paragraph_format.keep_with_next = True
            answer = " ".join(item["answer"])
            p = doc.add_paragraph()
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = True
            label = p.add_run("参考答案：")
            set_run_font(label, size=11, color=DARK_BLUE, bold=True)
            add_inline_runs(p, answer)


def build():
    sections = parse_source()
    total_questions = sum(len(s["questions"]) for s in sections)
    if total_questions != 150:
        raise ValueError(f"Expected 150 questions, found {total_questions}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(section)
    add_cover(doc)
    add_front_matter(doc, sections)
    add_qa_sections(doc, sections)

    props = doc.core_properties
    props.title = "RQ-VAE + HSTU 冷启动推荐项目：150 个面试问题与参考答案"
    props.subject = "推荐算法与机器学习工程面试准备"
    props.author = "Semantic-ID-Rec Project"
    props.keywords = "RQ-VAE, HSTU, Semantic ID, 冷启动, 推荐系统, 面试"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
